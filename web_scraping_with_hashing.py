from bs4 import BeautifulSoup
import requests
import docx
from hash_function import hashing
import os
import time

#You must enable the access for less secure apps in your google account.
#Otherwise you will get an error.
def send_notification_email(gmail:str,password:str):
    import smtplib

    server = smtplib.SMTP_SSL("smtp.gmail.com", 465)

    server.login(gmail, password)

    message = 'Just updated the docx file.'
    # your email, the other person's email( in that case you want to send the email to yourself)
    server.sendmail(gmail, gmail, message )

    server.quit()


while True:

    url = 'https://aws.amazon.com/new/?whats-new-content-all'

    response = requests.get(url)
    # print(response)


    is_added = False


    soup = BeautifulSoup(response.text, "html.parser")

    titles = soup.find_all("h5", {"class": "lb-txt-bold lb-txt-uppercase lb-txt-white lb-none-v-margin lb-h5 lb-title"})
    info = soup.find_all("div", {"class": "lb-txt-normal lb-txt-white lb-rtxt"})
    dates = soup.find_all("div", {"class": "lb-tiny-align-right lb-txt-bold lb-txt-none lb-txt-white lb-txt"})

    titles_string_list = [title.text.strip() for title in titles]
    info_string_list = [inf.text.strip() for inf in info]
    dates_string_list = [date.text.strip() for date in dates]


    all_in_one = [(titles_string_list[i],info_string_list[i],dates_string_list[i]) for i in range(len(dates_string_list))]
    hashed_objects = []


    for el in all_in_one:
        title = el[0]
        text = el[1]
        date = el[2]

        hash_value = hashing(title,text,date)
        hashed_objects.append(hash_value)




    if not os.path.isfile("D:/my_written_file.docx"):
        mydoc = docx.Document()

        for i in range(len(titles_string_list)):
            print(f"Title: {titles_string_list[i]}")
            mydoc.add_heading(titles_string_list[i])
            print(f"Text: {info_string_list[i]}")
            mydoc.add_paragraph(info_string_list[i])
            print(f"Date: {dates_string_list[i]}")
            mydoc.add_paragraph(dates_string_list[i])
            print(f"Hash value: {hashed_objects[i]}")
            print()
            mydoc.add_paragraph(f"")

        mydoc.save("D:/my_written_file.docx")

    else:
        mydoc = docx.Document("D:/my_written_file.docx")

        article = []

        articles = []

        for i, el in enumerate(mydoc.paragraphs):
            el = el.text

            article.append(el)
            if (i + 1) % 4 == 0:
                articles.append(article)
                article = []

        for article in articles:
            title, text, date, nothing = article

            hash_value = hashing(title, text, date)
            if hash_value not in hashed_objects:
                hashed_objects.append(hash_value)

                print(f"Title: {title}")
                mydoc.add_heading(title)
                print(f"Text: {text}")
                mydoc.add_paragraph(text)
                print(f"Date: {date}")
                mydoc.add_paragraph(date)
                print(f"Hash value: {hash_value}")
                print()
                mydoc.add_paragraph(f"")
                is_added = True


        if is_added:
            gmail = 'Write your gmail here'
            password = 'Write your password here'
            send_notification_email(gmail, password)

        try:
            mydoc.save("D:/my_written_file.docx")
        except:
            print('!!!Затвори word файла за да може да бъде ъпдейтнат'.upper())


    seconds = 6000
    time.sleep(seconds)
